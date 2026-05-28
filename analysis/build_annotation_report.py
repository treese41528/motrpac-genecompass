#!/usr/bin/env python3
"""
build_annotation_report.py — Generate collaborator report from report_data.json
Uses python-docx (no npm/node dependency).

Usage:
  python analysis/build_annotation_report.py
  # Output: reports/annotations/celltype_annotation_report.docx
"""

import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path("/depot/reese18/apps/motrpac-genecompass")
DATA_PATH = PROJECT_ROOT / "reports/annotations/report_data.json"
OUTPUT_PATH = PROJECT_ROOT / "reports/annotations/celltype_annotation_report.docx"

# Allow override via command line
if len(sys.argv) > 1:
    DATA_PATH = Path(sys.argv[1])
if len(sys.argv) > 2:
    OUTPUT_PATH = Path(sys.argv[2])


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None, header_color="1F4E79", alt_color="F2F7FB"):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Arial"
        set_cell_shading(cell, header_color)

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = "Arial"
            if r % 2 == 1:
                set_cell_shading(cell, alt_color)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_para(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return p


def add_rich_para(doc, segments):
    """Add paragraph with mixed formatting. segments: list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10.5)
        run.font.name = "Arial"
    return p


def build_report(data):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    # ===== TITLE PAGE =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Cell-Type Annotation for Deconvolution")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.font.name = "Arial"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MoTrPAC-GeneCompass Rat scRNA-seq Corpus")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Three-Method Consensus: PanglaoDB + ScType + LLM Reconciliation")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Tim Reese | Purdue University | {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()  # spacer

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)

    corpus = data["corpus"]
    fix = data["method_comparison"]["stellate_cell_fix"]
    add_rich_para(doc, [
        (f"We annotated {corpus['n_cells']/1e6:.1f}M cells across {corpus['n_samples']} samples from {corpus['n_motrpac_tissues']} MoTrPAC tissues using a three-method consensus pipeline. ", False),
        ("PanglaoDB marker scoring", True),
        (" provided broad coverage (178 cell types) but suffered from tissue-context mislabeling. ", False),
        ("ScType tissue-restricted scoring", True),
        (" corrected misassignments using curated tissue-specific cell type lists. ", False),
        ("LLM-based consensus reconciliation", True),
        (f" (Claude Haiku, 2,246 unique triples) resolved synonyms, recovered Unknown labels, and applied tissue-context corrections.", False),
    ])

    add_rich_para(doc, [
        ("Key result: ", True),
        (f"\"Pancreatic stellate cells\" — the #1 PanglaoDB cell type at {fix['panglao_total']/1000:.0f}K cells — was reduced by {fix['reduction_pct']}% to {fix['consensus_total']/1000:.0f}K cells "
         f"(retained only where biologically real). All {corpus['n_motrpac_tissues']} covered MoTrPAC tissues have biologically plausible compositions suitable for deconvolution.", False),
    ])

    # ===== 1. PIPELINE OVERVIEW =====
    add_heading(doc, "1. Three-Method Pipeline", level=1)

    add_heading(doc, "1.1 Method 1: PanglaoDB Marker Scoring", level=2)
    pdb = data.get("panglao_db", {})
    add_para(doc, f"PanglaoDB (v2020-03-27): {pdb.get('unique_cell_types', 178)} cell types, "
             f"{pdb.get('mapped_to_ensrnog', 7478)} markers mapped to ENSRNOG ({pdb.get('unmapped_symbols', 546)} unmapped). "
             f"Scoring via scanpy sc.tl.score_genes: mean z-score of marker gene expression vs expression-matched background. "
             f"Strengths: broad coverage, no tissue restriction needed. "
             f"Weakness: organ-specific labels for ubiquitous mesenchymal/immune markers (e.g., COL1A1/VIM/DCN scored as \"Pancreatic stellate cells\" everywhere).")

    add_heading(doc, "1.2 Method 2: ScType Tissue-Restricted Scoring", level=2)
    sdb = data.get("sctype_db", {})
    add_para(doc, f"ScTypeDB (Ianevski et al. 2022, Nat Commun): {sdb.get('total_entries', 268)} entries across "
             f"{sdb.get('n_tissues', 16)} tissues. {sdb.get('total_pos_mapped', 3633)} positive and "
             f"{sdb.get('total_neg_mapped', 17)} negative markers mapped to ENSRNOG. "
             f"Scoring on z-scaled expression of marker genes only (memory-efficient). "
             f"Tissue restriction: each sample scored against its matched ScTypeDB tissue + Immune system. "
             f"Threshold: clusters scoring below n_cells/4 labeled \"Unknown.\" "
             f"Negative markers had minimal practical impact (only 17 mapped, all in Immune system). "
             f"The tissue-restricted cell type lists are the mechanism that works.")

    add_heading(doc, "1.3 Method 3: LLM Consensus Reconciliation", level=2)
    llm = data.get("llm_corrections", {})
    add_para(doc, f"2,246 unique (PanglaoDB label, ScType label, tissue) triples were sent to Claude Haiku (batch_size=50, "
             f"45 batches, ~$0.10 total) with top-5 DE markers per cluster as context. The LLM applied five rules: "
             f"(1) synonym resolution, (2) tissue-context correction, (3) Unknown recovery from PanglaoDB, "
             f"(4) ScType preference for tissue-specific types, (5) DE marker tiebreaking. "
             f"{llm.get('total', 300)} labels were corrected. Results cached as consensus_label_map.json "
             f"and applied mechanically to all 16,790 clusters. 67 fallback clusters (0.4%) from a double-space "
             f"formatting mismatch in ScType's \"Natural killer  cells\" label.")

    # Source breakdown table
    add_heading(doc, "1.4 Consensus Source Breakdown", level=2)
    src_meanings = {
        "panglao": "PanglaoDB label confirmed by LLM",
        "sctype": "ScType label preferred by LLM",
        "merged": "Synonym resolved (both correct)",
        "corrected": "LLM overrode both methods",
        "fallback_panglao": "Triple not in cache (format mismatch)",
    }
    src_rows = []
    for src, info in data.get("consensus_sources", {}).items():
        src_rows.append([src, f"{info['n_clusters']:,}", f"{info['pct_clusters']:.1f}%",
                         f"{info['n_cells']:,}", src_meanings.get(src, "")])
    add_table(doc, ["Source", "Clusters", "% Clusters", "Cells", "Meaning"], src_rows,
              col_widths=[1.0, 0.8, 0.8, 1.0, 2.8])

    doc.add_page_break()

    # ===== 2. STELLATE CELL FIX =====
    add_heading(doc, "2. Tissue-Context Mislabeling Resolution", level=1)
    add_rich_para(doc, [
        ("The primary annotation artifact — ", False),
        (f"\"Pancreatic stellate cells\" at {fix['panglao_total']/1000:.0f}K cells (#1 PanglaoDB type)", True),
        (f" — was reduced to {fix['consensus_total']/1000:.0f}K cells ({fix['reduction_pct']}% reduction). ", False),
        ("Remaining stellate cells are biologically correct: hepatic stellate cells in liver (10.6K) and mesangial cells in kidney (106).", False),
    ])

    add_heading(doc, "2.1 Top LLM Corrections by Impact", level=2)
    corr_rows = []
    for corr in llm.get("top_corrections", [])[:15]:
        corr_rows.append([
            corr["panglao"], corr["consensus"], corr["tissue"],
            f"{corr.get('n_cells', 0)/1000:.0f}K",
            corr.get("rationale", "")[:55],
        ])
    add_table(doc, ["Original (PanglaoDB)", "Corrected To", "Tissue", "Cells", "Rationale"], corr_rows,
              col_widths=[1.6, 1.4, 0.8, 0.5, 2.1])

    doc.add_page_break()

    # ===== 3. PER-TISSUE RESULTS =====
    add_heading(doc, "3. Per-Tissue Cell Type Composition", level=1)
    add_para(doc, "Each tissue section shows final consensus cell types with cell counts, percentages, and top DE marker genes. "
             "Side-by-side method comparison is included for primary MoTrPAC tissues.")

    primary = ["heart", "kidney", "liver", "lung", "gastrocnemius"]
    secondary = ["cortex", "hippocampus", "hypothalamus", "blood RNA", "colon", "small intestine", "BAT", "WAT-SC", "spleen"]

    for idx, tissue in enumerate(primary):
        td = data.get("motrpac_tissues", {}).get(tissue, {})
        if not td or td.get("status") == "NO COVERAGE":
            continue

        add_heading(doc, f"3.{idx+1} {tissue.title()}", level=2)
        add_rich_para(doc, [
            (f"{td['n_cells']:,} cells", True),
            (f" | {td['n_samples']} samples | {td['n_studies']} studies | {td['n_cell_types']} cell types", False),
        ])

        # Method comparison table
        mc = data.get("tissue_method_comparison", {}).get(tissue, {})
        if mc:
            add_para(doc, "Method Comparison (Top 5):", bold=True, size=10)
            mc_rows = []
            for i in range(min(5, max(len(mc.get("panglao_top5", [])), len(mc.get("sctype_top5", [])), len(mc.get("consensus_top5", []))))):
                pg = mc.get("panglao_top5", [{}])[i] if i < len(mc.get("panglao_top5", [])) else {}
                st = mc.get("sctype_top5", [{}])[i] if i < len(mc.get("sctype_top5", [])) else {}
                cn = mc.get("consensus_top5", [{}])[i] if i < len(mc.get("consensus_top5", [])) else {}
                mc_rows.append([
                    f"{pg.get('type', '')} ({pg.get('cells', 0)/1000:.0f}K)" if pg.get("type") else "",
                    f"{st.get('type', '')} ({st.get('cells', 0)/1000:.0f}K)" if st.get("type") else "",
                    f"{cn.get('type', '')} ({cn.get('cells', 0)/1000:.0f}K)" if cn.get("type") else "",
                ])
            add_table(doc, ["PanglaoDB", "ScType", "Consensus (Final)"], mc_rows,
                      col_widths=[2.2, 2.2, 2.2])

        # Cell type composition table
        add_para(doc, "Final Consensus Composition:", bold=True, size=10)
        de_markers = data.get("de_markers_by_tissue", {}).get(tissue, {})
        comp_rows = []
        for ct in td.get("cell_types", [])[:12]:
            markers = de_markers.get(ct["cell_type"], [])
            marker_str = ", ".join(m["symbol"] for m in markers[:3]) if markers else ""
            comp_rows.append([
                ct["cell_type"], f"{ct['n_cells']:,}", f"{ct['pct']:.1f}%", marker_str,
            ])
        add_table(doc, ["Cell Type", "Cells", "%", "Top DE Markers"], comp_rows,
                  col_widths=[2.2, 1.0, 0.6, 2.6])

        doc.add_page_break()

    # Secondary tissues (compact)
    add_heading(doc, f"3.{len(primary)+1} Additional MoTrPAC Tissues", level=2)
    for tissue in secondary:
        td = data.get("motrpac_tissues", {}).get(tissue, {})
        if not td or td.get("status") == "NO COVERAGE":
            continue
        top5 = ", ".join(f"{ct['cell_type']} ({ct['pct']}%)" for ct in td.get("cell_types", [])[:5])
        add_rich_para(doc, [
            (f"{tissue}: ", True),
            (f"{td['n_cells']:,} cells, {td['n_cell_types']} types, {td['n_samples']} samples. Top: {top5}", False),
        ])

    doc.add_page_break()

    # ===== 4. DECONVOLUTION READINESS =====
    add_heading(doc, "4. Deconvolution Readiness", level=1)
    n_ready = sum(1 for v in data.get("deconvolution_readiness", {}).values() if v.get("ready"))
    add_para(doc, f"All {n_ready} covered MoTrPAC tissues meet the minimum threshold (50+ cells per cell type "
             f"for at least 3 cell types) for deconvolution reference panel construction.")

    ready_rows = []
    for tissue in primary + secondary:
        rd = data.get("deconvolution_readiness", {}).get(tissue, {})
        if not rd:
            continue
        td = data.get("motrpac_tissues", {}).get(tissue, {})
        ready_rows.append([
            tissue,
            str(rd.get("n_cell_types_total", 0)),
            str(rd.get("n_cell_types_above_threshold", 0)),
            f"{td.get('n_cells', 0):,}",
            "READY" if rd.get("ready") else "NOT READY",
        ])
    add_table(doc, ["Tissue", "Cell Types", "Above 50-cell", "Total Cells", "Status"], ready_rows,
              col_widths=[1.3, 0.8, 1.0, 1.2, 1.0])

    # Coverage gaps
    add_heading(doc, "4.1 Coverage Gaps", level=2)
    gaps = data.get("coverage_gaps", {})
    add_rich_para(doc, [
        ("Missing tissues (no scRNA-seq data in corpus): ", True),
        (", ".join(gaps.get("missing_tissues", [])), False),
    ])
    thin_str = "; ".join(f"{t['tissue']} ({t['n_cells']:,} cells, {t['n_samples']} samples)"
                          for t in gaps.get("thin_tissues", []))
    add_rich_para(doc, [
        ("Thin tissues (<50K cells): ", True),
        (thin_str, False),
    ])

    # Tool compatibility
    add_heading(doc, "4.2 Tool Compatibility", level=2)
    for tool, info in data.get("deconvolution_tools", {}).items():
        add_rich_para(doc, [
            (f"{tool}: ", True),
            (info.get("status", ""), False),
        ])
        if info.get("note"):
            add_para(doc, f"  Note: {info['note']}", italic=True, size=9)

    doc.add_page_break()

    # ===== 5. PIPELINE PARAMETERS =====
    add_heading(doc, "5. Pipeline Parameters", level=1)
    pp = data.get("pipeline_parameters", {})
    param_rows = [
        ["Normalization", pp.get("normalization", ""), "Standard scanpy"],
        ["HVG selection", pp.get("hvg_selection", ""), "Top 2000 most variable"],
        ["Leiden resolution", str(pp.get("leiden_resolution", "")), "Moderate granularity"],
        ["PCA components", pp.get("pca_components", ""), "Adapted to sample size"],
        ["ScType threshold", pp.get("sctype_unknown_threshold", ""), "Below → Unknown"],
        ["DE method", pp.get("de_method", ""), "Wilcoxon rank-sum"],
        ["Consensus LLM", pp.get("consensus_llm_model", ""), "Claude Haiku 4.5"],
        ["Batch size", str(pp.get("consensus_batch_size", "")), "50 triples/call"],
    ]
    add_table(doc, ["Parameter", "Value", "Rationale"], param_rows, col_widths=[1.5, 2.8, 2.1])

    # ===== 6. SAMPLE METADATA =====
    add_heading(doc, "6. Sample Metadata", level=1)
    sm = data.get("sample_metadata", {})
    add_para(doc, f"{sm.get('n_samples', 0)} samples. {sm.get('pct_with_tissue', 0)}% have tissue metadata, "
             f"{sm.get('pct_with_condition', 0)}% have condition metadata.")

    strain_rows = [[s, str(n)] for s, n in list(sm.get("strain_distribution", {}).items())[:8]]
    add_table(doc, ["Strain", "Samples"], strain_rows, col_widths=[3.0, 1.0])

    # ===== 7. ACTION ITEMS =====
    add_heading(doc, "7. Action Items", level=1)

    add_heading(doc, "7.1 For Domain Experts (Dr. Gavin, Fei Xue)", level=2)
    for item in [
        "Review per-tissue cell type compositions in Section 3. Are the top cell types biologically plausible?",
        "Check DE markers for top cell types. Do marker genes match expected identities?",
        "Assess deconvolution readiness. Are there sufficient cell types per tissue?",
        "Identify any cell types that should be present but are missing.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "7.2 For Yang He", level=2)
    for item in [
        "Spot-check 2-3 samples per tissue with UMAPs.",
        "Cross-reference consensus compositions with published rat single-cell atlases.",
        "Evaluate whether Leiden resolution 0.8 is appropriate per tissue.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "7.3 For Fei Xue (Deconvolution)", level=2)
    for item in [
        "Review deconvolution readiness in Section 4.",
        "Determine UniCell approach: UCDBase fine-tuning to rat or UCDSelect with custom references?",
        "Plan scDEAL KNN embedding transfer: which tissues × timepoints to prioritize?",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ===== 8. REPRODUCIBILITY =====
    add_heading(doc, "8. Reproducibility", level=1)
    add_para(doc, "All scripts, databases, and outputs are in the project repository:")
    for item in [
        "analysis/annotate_celltypes.py — PanglaoDB pipeline",
        "analysis/annotate_sctype.py — ScType pipeline",
        "analysis/consensus_annotations.py — LLM consensus merger",
        "analysis/extract_report_data.py — Data extraction for this report",
        "data/training/cell_annotations_consensus/ — Final consensus annotations",
        "reports/annotations/report_data.json — Complete data (179 KB)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MoTrPAC-GeneCompass | NIH Grant Period: December 2025 – November 2026 | Purdue University")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Report saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    print(f"Loading data from {DATA_PATH}...")
    with open(DATA_PATH) as f:
        data = json.load(f)
    build_report(data)